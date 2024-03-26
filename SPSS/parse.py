import fitz  # PyMuPDF
import re
from docx import Document

# Open the PDF file
pdf_document = fitz.open('palliativeeedata.pdf')

# Prepare to extract text
text = ""
for page_num in range(len(pdf_document)):
    page = pdf_document[page_num]
    text += page.get_text()

# Close the PDF after extraction
pdf_document.close()

# Define the regex pattern for extracting table data
pattern = re.compile(
    r"ONEWAY\s+([\w\s]+)\s+BY\s+([\w\s]+)\s+([\s\S]*?)(?=(?:ONEWAY|$))"
)

# Find matches using the pattern
matches = pattern.findall(text)

# Function to parse table data from the extracted section data
def parse_table_data(section_data):
    rows = section_data.strip().split('\n')
    table_data = [row.split() for row in rows if row.strip()]
    return table_data

# Load the existing Word document
doc = Document('pa_table.docx')

# Function to fill a table in the Word document with the given data
def fill_table(doc, table_data, headers):
    table = doc.add_table(rows=len(table_data), cols=len(headers))
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
    for i, row_data in enumerate(table_data, start=1):
        for j, cell_data in enumerate(row_data):
            table.cell(i, j).text = str(cell_data)

# Iterate over matches and fill tables in the Word document
for match in matches:
    category, subcategory, table_data = match[0], match[1], match[2]
    headers = ["Variable", "N", "Mean", "Std. Deviation", "Std. Error", "95% Confidence Interval for Mean", 
               "Lower Bound", "Upper Bound", "ANOVA", "Sum of Squares", "df", "Mean Square", "F", "Sig."]
    table_data = parse_table_data(table_data)
    
    # Fill table in Word document
    doc.add_paragraph(f"{category} by {subcategory}")
    fill_table(doc, table_data, headers)

# Save the updated document
updated_doc_path = '/mnt/data/Complete_Palliative_Tables.docx'
doc.save(updated_doc_path)

print(f"Updated document saved at: {updated_doc_path}")
